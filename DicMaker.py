# -*- coding: utf-8 -*-
import os
import shutil
from os.path import *
from random import randint

import fitz
from PIL import Image as PILImage, ImageFont, ImageDraw
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import PageTemplate, Frame
from reportlab.platypus import SimpleDocTemplate, Image as rep_Image, PageBreak
from win32com.client import Dispatch

from Searcher import Searcher


class DicMaker:
    def __init__(self, wordlist, path, gui=None, qpb=None, is_header=False, title=None, pre_load=None):
        self.Searcher = Searcher(pre_load)
        self.doc = Document()
        self.path = path
        self.is_header = is_header
        self.index = randint(100000, 999999)
        try:
            os.mkdir("{}/Temps".format(dirname(abspath(__file__))))
        except:
            pass
        self.temp_path = dirname(abspath(__file__)) + "/Temps/{}.pdf".format(self.index)
        self.gui = gui
        self.page_num = 0
        dl = 1 / len(wordlist) * 100
        exact = 0
        self.doc.save(self.path)
        if self.is_header:
            self.load_word()
        self.initals = []
        self.headers = []
        self.special_section = []
        begin, end = wordlist[0], ""
        initial, page = "", 1
        for i in range(0, len(wordlist)):
            exact += dl
            if qpb is not None:
                qpb.setValue(int(exact))
            if self.gui is not None:
                self.gui()
            firstcheck, secondcheck = False, False
            flag = False
            word = wordlist[i]
            if word[0] != initial:
                self.addinitial(word[0])
                initial = word[0]
                section = self.doc.add_section(start_type=0)
                section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
                flag = True
            if self.is_header:
                if self.checkpage() != page:
                    page = self.checkpage()
                    firstcheck = True
            self.addword(word)
            if self.gui is not None:
                self.gui()
            if self.is_header:
                if self.checkpage() != page:
                    page = self.checkpage()
                    secondcheck = True
            if self.is_header:
                if firstcheck:
                    end = wordlist[i - 1]
                    self.headers.append((begin, end))
                    if flag:
                        begin = word
                else:
                    if secondcheck:
                        end = word
                        self.headers.append((begin, end))
                        begin = word
                if flag:
                    self.initals.append((word[0], page))
        if self.is_header:
            self.headers.append((begin, wordlist[-1]))
        self.save()
        if self.is_header:
            if qpb is not None:
                qpb.setValue(0)
            if title is not None:
                title.setText("页码添加中")
            self.addcontent()
            self.save()
            if self.gui is not None:
                self.gui()
            self.word2pdf()
            if qpb is not None:
                qpb.setValue(25)
            if self.gui is not None:
                self.gui()
            self.pdf2jpeg()
            if qpb is not None:
                qpb.setValue(50)
            if self.gui is not None:
                self.gui()
            self.add_header()
            if qpb is not None:
                qpb.setValue(75)
            if self.gui is not None:
                self.gui()
            self.jpeg2pdf()
            if title is not None:
                title.setText("页码添加完成，正在清理缓存")
            self.word_File.Close()
            self.Word_operator.Quit()
        shutil.rmtree("{}/Temps".format(dirname(abspath(__file__))))

    def addword(self, word):
        pronounce, meanings, state = self.Searcher.search(word)
        if state == -1:
            return
        p = self.doc.add_paragraph()
        run = p.add_run(word)
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True
        run = p.add_run(
            " {}".format(pronounce.translate(str.maketrans('[', '/', '')).translate(str.maketrans(']', '/', ''))))
        run.font.name = "Arial"
        run.font.size = Pt(14)
        for mean in meanings:
            if mean.string is None:
                continue
            try:
                speech, cn = mean.string.split(".")
            except:
                cn = mean.string
                run = p.add_run(" •{}".format(cn))
                run.font.name = u"等线"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), u"等线")
                run.font.size = Pt(11)
                continue
            run = p.add_run(" •{}.".format(speech))
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run = p.add_run(cn)
            run.font.name = u"等线"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u"等线")
            run.font.size = Pt(11)

    def checkpage(self):
        flag = True
        while (flag):
            if self.gui is not None:
                self.gui()
            try:
                self.save()
                flag = False
                return self.word_File.ComputeStatistics(2)
            except:
                pass

    def addinitial(self, letter):
        seclen = len(self.doc.sections)
        if seclen == 1:
            self.special_section.append(0)
        else:
            self.special_section.append(seclen)
            self.doc.add_section(start_type=2)._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
        p = self.doc.add_paragraph()
        run = p.add_run("{}".format(letter.upper()))
        run.font.name = "Arial"
        run.font.size = Pt(20)
        run.font.bold = True
        p.alignment = 1

    def load_word(self):
        self.Word_operator = Dispatch('Word.Application')
        self.Word_operator.Visible = False
        self.Word_operator.DisplayAlerts = 0
        self.word_File = self.Word_operator.Documents.Open(self.path)

    def save(self):
        if self.is_header:
            self.word_File.Close()
            self.doc.save(self.path)
            self.word_File = self.Word_operator.Documents.Open(self.path)
        else:
            self.doc.save(self.path)

    def addcontent(self):
        self.content = self.doc.paragraphs[0].insert_paragraph_before()
        run = self.content.add_run("Content\n")
        run.font.name = "Consolas"
        run.font.size = Pt(20)
        run.font.bold = True
        for inital in self.initals:
            if self.gui is not None:
                self.gui()
            run = self.content.add_run("{}\t……\t{}\n".format(inital[0].upper(), inital[1]))
            run.font.name = "Consolas"
            run.font.size = Pt(16)
        self.content.add_run().add_break(7)

    def word2pdf(self):
        self.save()
        self.word_File.SaveAs(self.temp_path, FileFormat=17)

    def pdf2jpeg(self):
        doc = fitz.open(self.temp_path)
        rotate, zoom_x, zoom_y = int(0), 2.0, 2.0
        trans = fitz.Matrix(zoom_x, zoom_y).preRotate(rotate)
        self.page_num = doc.pageCount
        for i in range(self.page_num):
            if self.gui is not None:
                self.gui()
            page = doc[i]
            pm = page.getPixmap(matrix=trans, alpha=False)
            pm.writeImage("{}_{}.jpeg".format(self.temp_path[:-4], str(i)))

    def add_header(self):
        font_path = "C:/Windows/Fonts/consola.ttf"
        font = ImageFont.truetype(font_path, 18)
        for i in range(1, self.page_num):
            if self.gui is not None:
                self.gui()
            header = "{} —— {}".format(self.headers[i - 1][0], self.headers[i - 1][1])
            image = PILImage.open("{}_{}.jpeg".format(self.temp_path[:-4], str(i)))
            draw = ImageDraw.Draw(image)
            width = draw.textsize(header, font=font)[0]
            draw.text(((image.size[0] - width) / 2, 50), header, font=font, fill='#808080')
            width = draw.textsize(str(i), font=font)[0]
            draw.text(((image.size[0] - width) / 2, image.size[1] - 65), str(i), font=font, fill='#808080')
            image.save("{}_{}.jpeg".format(self.temp_path[:-4], str(i)))

    def jpeg2pdf(self):
        file_name = "{}.pdf".format(self.path[:-5])
        width, height = landscape(A4)
        doc = SimpleDocTemplate(file_name, pagesize=(height, width))
        frame = Frame(0, 0, height, width, 0, 0, 0, 0)
        doc.addPageTemplates([PageTemplate(id="Later", frames=frame)])
        Story = []
        pictures = ["{}_{}.jpeg".format(self.temp_path[:-4], str(i)) for i in range(0, self.page_num)]
        for pic in pictures:
            if self.gui is not None:
                self.gui()
            img = rep_Image(pic, height, width)
            Story.append(img)
            Story.append(PageBreak())
        doc.build(Story)
